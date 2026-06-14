"""
CV Parser Service
Extracts text from PDF/DOCX files and uses Google Gemini to parse structured data.
"""
import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, Optional

from google import genai
from google.genai import types

from app.core.config import settings

logger = logging.getLogger(__name__)


def _get_client() -> genai.Client:
    if not settings.GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY is not set in environment variables")
    return genai.Client(api_key=settings.GEMINI_API_KEY)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF (fitz) or fallback to PyPDF2."""
    text = ""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except ImportError:
        logger.warning("PyMuPDF not available, falling back to PyPDF2")

    try:
        import PyPDF2
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise RuntimeError(f"Could not extract text from PDF: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise RuntimeError(f"Could not extract text from DOCX: {e}")


def extract_text_from_file(file_path: str) -> str:
    """Auto-detect file type and extract text."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _clean_json_response(raw: str) -> str:
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return raw.strip()


async def parse_cv(file_path: str) -> Dict[str, Any]:
    """
    Parse CV file using Google Gemini.

    Returns a dict with:
        name, email, phone, skills, education, work_experience, certifications
    """
    client = _get_client()

    raw_text = extract_text_from_file(file_path)
    if not raw_text:
        raise ValueError("CV file appears to be empty or unreadable")

    raw_text_truncated = raw_text[:15000]

    prompt = f"""
You are an expert CV/Resume parser. Analyze the following CV text and extract structured information.

CV TEXT:
---
{raw_text_truncated}
---

Return ONLY a valid JSON object (no markdown, no explanation) with this exact structure:
{{
  "name": "Full name of the candidate",
  "email": "email@example.com",
  "phone": "phone number or null",
  "skills": ["skill1", "skill2", "skill3"],
  "education": [
    {{
      "institution": "University name",
      "degree": "Bachelor/Master/PhD/etc",
      "field": "Field of study",
      "year": "graduation year or duration"
    }}
  ],
  "work_experience": [
    {{
      "company": "Company name",
      "title": "Job title",
      "duration": "e.g. 2020-2023 or 3 years",
      "responsibilities": ["responsibility 1", "responsibility 2"]
    }}
  ],
  "certifications": ["Certification 1", "Certification 2"],
  "summary": "Brief professional summary if present"
}}

Rules:
- Extract ALL skills mentioned (technical, soft, tools, languages, frameworks)
- List ALL work experiences in reverse chronological order
- If a field is not found, use null or empty array []
- Return ONLY the JSON object, no other text
"""

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.1,
                max_output_tokens=4096,
            ),
        )
        raw_json = _clean_json_response(response.text)
        parsed = json.loads(raw_json)

        parsed.setdefault("name", "Unknown")
        parsed.setdefault("email", "")
        parsed.setdefault("phone", None)
        parsed.setdefault("skills", [])
        parsed.setdefault("education", [])
        parsed.setdefault("work_experience", [])
        parsed.setdefault("certifications", [])
        parsed.setdefault("summary", "")

        return parsed

    except json.JSONDecodeError as e:
        logger.error(f"Gemini returned invalid JSON: {e}")
        return {
            "name": _extract_name_fallback(raw_text),
            "email": _extract_email_fallback(raw_text),
            "phone": _extract_phone_fallback(raw_text),
            "skills": [],
            "education": [],
            "work_experience": [],
            "certifications": [],
            "summary": raw_text[:500],
        }
    except Exception as e:
        logger.error(f"Gemini CV parsing failed: {e}")
        raise RuntimeError(f"CV parsing failed: {e}")


def _extract_email_fallback(text: str) -> str:
    match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else ""


def _extract_phone_fallback(text: str) -> Optional[str]:
    match = re.search(r"(\+?\d[\d\s\-\(\)]{7,15}\d)", text)
    return match.group(0).strip() if match else None


def _extract_name_fallback(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return lines[0] if lines else "Unknown"
